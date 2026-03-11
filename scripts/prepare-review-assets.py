#!/usr/bin/env python3
"""
prepare-review-assets.py — 学术委员会审前一键资产预处理

合并 paper-asset-extractor.py + preprocess-paper-assets.py 的功能。

输入: 论文文件 (.md / .pdf / .docx / .tex)
输出:
  review-board/cache/paper-figure-pack.json   — 结构化数据
  review-board/cache/paper-figure-pack.md     — 专家可读文本

流程:
  1. 解析论文 → 章节 + 表格 + 图片文件
  2. 分类图片: 数据图 → 提取生成代码; 概念图 → 标记待描述
  3. (可选) --fill-descriptions 填充视觉描述
  4. 输出 JSON + Markdown

用法:
  # 基础用法 (Markdown 论文)
  python3 prepare-review-assets.py ./paper.md

  # PDF / Word / LaTeX
  python3 prepare-review-assets.py ./paper.pdf
  python3 prepare-review-assets.py ./main.tex --tex-root .

  # 填充概念图的视觉描述
  python3 prepare-review-assets.py ./paper.md \\
    --fill-descriptions cache/figure-descriptions.json

  # 指定额外的图片目录 (覆盖自动检测)
  python3 prepare-review-assets.py ./paper.pdf --figures-dir ./figures/

  # 注册新图 (追加到 registry)
  python3 prepare-review-assets.py --register new-fig.png \\
    --paper-id "Figure 4" --fig-type data \\
    --generator "generate_ict_figures.py::generate_figure4"
"""

import argparse
import json
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent
REVIEW_DIR = SCRIPT_DIR.parent
PROJECT_DIR = REVIEW_DIR.parent
DEFAULT_FIGURES_DIR = PROJECT_DIR / "figures"
CACHE_DIR = REVIEW_DIR / "cache"
REGISTRY_PATH = CACHE_DIR / "figure-registry.json"


# ═══════════════════════════════════════════════════════════════════════════════
# Part 1: 论文解析 (PDF / Word / LaTeX / Markdown)
# ═══════════════════════════════════════════════════════════════════════════════

def make_output_dir(output_dir: str) -> Path:
    p = Path(output_dir)
    p.mkdir(parents=True, exist_ok=True)
    return p


# ── Markdown ──────────────────────────────────────────────────────────────────

def extract_markdown(input_path: str, figures_dir: Path) -> dict:
    """从 Markdown 提取章节、表格、图片引用"""
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()
    lines = content.split("\n")

    # 章节: # / ## / ###
    sections = []
    current_section = None
    current_content = []
    sec_counter = 0

    def flush():
        nonlocal current_section
        if current_section is not None:
            current_section["content"] = "\n".join(current_content).strip()
            sections.append(current_section)
        current_content.clear()

    for line in lines:
        heading_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if heading_match:
            flush()
            sec_counter += 1
            current_section = {
                "id": f"sec-{sec_counter}",
                "title": heading_match.group(2).strip(),
                "level": len(heading_match.group(1)),
                "content": "",
            }
        else:
            current_content.append(line)
    flush()

    # 表格: **Table N: ...**  后跟 | 行
    tables = []
    tab_pattern = re.compile(r"\*\*Table\s+(\d+):\s*(.+?)\*\*")
    i = 0
    while i < len(lines):
        m = tab_pattern.match(lines[i])
        if m:
            tab_num, tab_title = m.group(1), m.group(2).strip()
            md_lines = []
            j = i + 1
            while j < len(lines):
                if lines[j].strip().startswith("|"):
                    md_lines.append(lines[j])
                    j += 1
                elif not lines[j].strip() and md_lines:
                    break
                elif not lines[j].strip():
                    j += 1
                else:
                    break
            ctx_start = max(0, i - 2)
            tables.append({
                "id": f"Table {tab_num}",
                "title": tab_title,
                "markdown": "\n".join(md_lines) if md_lines else "(inline in text)",
                "context": "\n".join(lines[ctx_start:i]).strip(),
                "page": None,
            })
        i += 1

    # 图片: ![...](path)
    figures = []
    fig_counter = 0
    img_pattern = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
    for line in lines:
        for m in img_pattern.finditer(line):
            fig_counter += 1
            caption = m.group(1)
            img_rel = m.group(2)
            # 解析路径
            img_path = (Path(input_path).parent / img_rel).resolve()
            if not img_path.exists():
                img_path = figures_dir / Path(img_rel).name
            figures.append({
                "id": f"fig-{fig_counter}",
                "caption": caption,
                "section": sections[-1]["id"] if sections else "",
                "source_path": img_rel,
                "extracted_image_path": str(img_path) if img_path.exists() else "",
                "page": None,
            })

    # 也扫描 figures_dir 中 paper 未引用的图
    if figures_dir.exists():
        referenced = {Path(f["source_path"]).name for f in figures}
        for img_file in sorted(figures_dir.glob("*.png")) + sorted(figures_dir.glob("*.jpg")):
            if img_file.name not in referenced and not img_file.name.startswith("."):
                fig_counter += 1
                figures.append({
                    "id": f"fig-{fig_counter}",
                    "caption": "",
                    "section": "",
                    "source_path": str(img_file),
                    "extracted_image_path": str(img_file),
                    "page": None,
                })

    return {"sections": sections, "tables": tables, "figures": figures, "text": content}


# ── PDF ───────────────────────────────────────────────────────────────────────

def extract_pdf(input_path: str, output_dir: Path) -> dict:
    try:
        import fitz
    except ImportError:
        sys.exit("❌ 缺少 PyMuPDF，请运行: pip install PyMuPDF")

    doc = fitz.open(input_path)

    # 字体大小分布 → 正文字体
    font_size_chars: dict[float, int] = {}
    for page in doc:
        for blk in page.get_text("dict")["blocks"]:
            if blk["type"] != 0:
                continue
            for line in blk.get("lines", []):
                for span in line.get("spans", []):
                    sz = round(span["size"], 1)
                    font_size_chars[sz] = font_size_chars.get(sz, 0) + len(span.get("text", ""))
    body_size = max(font_size_chars, key=font_size_chars.get) if font_size_chars else 10.0

    sections, tables, all_text_lines = [], [], []
    current_section, current_content = None, []
    sec_counter = tab_counter = 0

    def flush_sec():
        nonlocal current_section
        if current_section is not None:
            current_section["content"] = "\n".join(current_content).strip()
            sections.append(current_section)
        current_content.clear()

    for page_num, page in enumerate(doc, 1):
        for blk in page.get_text("dict")["blocks"]:
            if blk["type"] != 0:
                continue
            for line in blk.get("lines", []):
                line_text = "".join(s["text"] for s in line.get("spans", [])).strip()
                if not line_text:
                    continue
                all_text_lines.append(line_text)

                is_heading, heading_level = False, 1
                for span in line.get("spans", []):
                    size = round(span["size"], 1)
                    is_bold = bool(span.get("flags", 0) & (1 << 4))
                    if size >= body_size + 3:
                        is_heading = True
                        heading_level = 1 if size >= body_size + 6 else (2 if size >= body_size + 4 else 3)
                        break
                    elif is_bold and size >= body_size + 1:
                        is_heading, heading_level = True, 2
                        break

                if is_heading:
                    flush_sec()
                    sec_counter += 1
                    current_section = {"id": f"sec-{sec_counter}", "title": line_text, "level": heading_level, "content": ""}
                else:
                    current_content.append(line_text)

        # 表格检测
        try:
            for tbl in page.find_tables().tables:
                tab_counter += 1
                extracted = tbl.extract()
                md_rows = []
                for row_idx, row in enumerate(extracted):
                    cells = [str(c or "").strip().replace("\n", " ") for c in row]
                    md_rows.append("| " + " | ".join(cells) + " |")
                    if row_idx == 0:
                        md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
                tables.append({
                    "id": f"tab-{tab_counter}", "title": "", "markdown": "\n".join(md_rows),
                    "context": "", "page": page_num,
                })
        except AttributeError:
            pass

    flush_sec()

    # 图片提取
    figures = []
    fig_counter = 0
    seen_xrefs = set()
    for page_num, page in enumerate(doc, 1):
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            fig_counter += 1
            fig_id = f"fig-{fig_counter}"
            try:
                base_image = doc.extract_image(xref)
                ext = base_image["ext"]
                if ext == "jpeg":
                    ext = "jpg"
                save_path = output_dir / f"{fig_id}.{ext}"
                with open(save_path, "wb") as f:
                    f.write(base_image["image"])
                figures.append({
                    "id": fig_id, "caption": "", "section": sections[-1]["id"] if sections else "",
                    "source_path": "", "extracted_image_path": str(save_path), "page": page_num,
                })
            except Exception as e:
                print(f"⚠️  跳过图片 xref={xref} (p{page_num}): {e}")

    doc.close()
    return {"sections": sections, "tables": tables, "figures": figures, "text": "\n".join(all_text_lines)}


# ── Word (.docx) ──────────────────────────────────────────────────────────────

def extract_docx(input_path: str, output_dir: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        sys.exit("❌ 缺少 python-docx，请运行: pip install python-docx")

    doc = Document(input_path)
    sections, current_section, current_content = [], None, []
    sec_counter = 0
    all_text = []

    def flush():
        nonlocal current_section
        if current_section is not None:
            current_section["content"] = "\n".join(current_content).strip()
            sections.append(current_section)
        current_content.clear()

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if text:
            all_text.append(text)
        if style.startswith("Heading"):
            flush()
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            sec_counter += 1
            current_section = {"id": f"sec-{sec_counter}", "title": text, "level": level, "content": ""}
        elif text:
            current_content.append(text)
    flush()

    # 表格
    tables = []
    for tab_idx, table in enumerate(doc.tables, 1):
        md_rows = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            md_rows.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        tables.append({
            "id": f"tab-{tab_idx}", "title": "", "markdown": "\n".join(md_rows),
            "context": "", "page": None,
        })

    # 图片
    figures = []
    fig_counter = 0
    image_parts = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_parts[rel_id] = rel.target_part

    for para in doc.paragraphs:
        rids = re.findall(r'r:embed="(rId\d+)"', para._element.xml)
        for rid in rids:
            if rid in image_parts:
                fig_counter += 1
                img_part = image_parts.pop(rid)
                ext = img_part.content_type.split("/")[-1]
                if ext == "jpeg":
                    ext = "jpg"
                save_path = output_dir / f"fig-{fig_counter}.{ext}"
                with open(save_path, "wb") as f:
                    f.write(img_part.blob)
                figures.append({
                    "id": f"fig-{fig_counter}", "caption": "", "section": sections[-1]["id"] if sections else "",
                    "source_path": "", "extracted_image_path": str(save_path), "page": None,
                })

    return {"sections": sections, "tables": tables, "figures": figures, "text": "\n".join(all_text)}


# ── LaTeX ─────────────────────────────────────────────────────────────────────

def extract_tex(input_path: str, output_dir: Path, tex_root: str = None) -> dict:
    root = Path(tex_root).expanduser() if tex_root else Path(input_path).parent

    with open(input_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    content_nc = re.sub(r"(?<!\\)%[^\n]*", "", content)

    # 章节
    sec_pattern = re.compile(r"\\((?:sub)*section)\*?\{([^}]+)\}", re.DOTALL)
    sec_matches = list(sec_pattern.finditer(content_nc))
    level_map = {"section": 1, "subsection": 2, "subsubsection": 3}
    sections = []
    sec_counter = 0

    for idx, m in enumerate(sec_matches):
        sec_counter += 1
        start = m.end()
        end = sec_matches[idx + 1].start() if idx + 1 < len(sec_matches) else len(content_nc)
        raw = content_nc[start:end]
        text = _clean_latex(raw)
        sections.append({
            "id": f"sec-{sec_counter}", "title": m.group(2).strip(),
            "level": level_map.get(m.group(1), 1), "content": text.strip(),
        })

    # 表格
    table_pattern = re.compile(r"\\begin\{table[*]?\}(.*?)\\end\{table[*]?\}", re.DOTALL)
    tabular_pattern = re.compile(r"\\begin\{tabular\}.*?\\end\{tabular\}", re.DOTALL)
    caption_pattern = re.compile(r"\\caption\{([^}]+)\}")
    tables = []
    tab_counter = 0

    for m in table_pattern.finditer(content_nc):
        tab_counter += 1
        body = m.group(1)
        cap_m = caption_pattern.search(body)
        tab_m = tabular_pattern.search(body)
        tables.append({
            "id": f"tab-{tab_counter}",
            "title": cap_m.group(1).strip() if cap_m else "",
            "markdown": _tabular_to_markdown(tab_m.group(0)) if tab_m else _clean_latex(body),
            "context": "", "page": None,
        })

    # 图片
    fig_pattern = re.compile(r"\\begin\{figure[*]?\}(.*?)\\end\{figure[*]?\}", re.DOTALL)
    incl_pattern = re.compile(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}")
    figures = []
    fig_counter = 0

    for m in fig_pattern.finditer(content_nc):
        body = m.group(1)
        cap_m = caption_pattern.search(body)
        caption = cap_m.group(1).strip() if cap_m else ""
        for img_m in incl_pattern.finditer(body):
            fig_counter += 1
            rel_path = img_m.group(1).strip()
            abs_path = _resolve_tex_image(root, rel_path)
            extracted = ""
            if abs_path and abs_path.exists():
                dest = output_dir / f"fig-{fig_counter}{abs_path.suffix}"
                shutil.copy2(str(abs_path), str(dest))
                extracted = str(dest)
            figures.append({
                "id": f"fig-{fig_counter}", "caption": caption,
                "section": sections[-1]["id"] if sections else "",
                "source_path": rel_path, "extracted_image_path": extracted, "page": None,
            })

    return {"sections": sections, "tables": tables, "figures": figures, "text": content_nc}


def _clean_latex(text: str) -> str:
    text = re.sub(r"\\begin\{[^}]+\}.*?\\end\{[^}]+\}", " ", text, flags=re.DOTALL)
    text = re.sub(r"\\[a-zA-Z]+\*?\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?", " ", text)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _tabular_to_markdown(tabular_str: str) -> str:
    body = re.sub(r"\\begin\{tabular\}[^}]*\}?", "", tabular_str)
    body = re.sub(r"\\end\{tabular\}", "", body)
    body = re.sub(r"\\hline|\\cline\{[^}]*\}", "", body)
    rows = []
    for raw_row in re.split(r"\\\\", body):
        raw_row = raw_row.strip()
        if not raw_row:
            continue
        cells = [_clean_latex(c).strip() for c in raw_row.split("&")]
        if any(c for c in cells):
            rows.append(cells)
    if not rows:
        return ""
    max_cols = max(len(r) for r in rows)
    md_lines = []
    for row_idx, row in enumerate(rows):
        padded = row + [""] * (max_cols - len(row))
        md_lines.append("| " + " | ".join(padded) + " |")
        if row_idx == 0:
            md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    return "\n".join(md_lines)


def _resolve_tex_image(root: Path, rel_path: str) -> Path | None:
    candidate = root / rel_path
    if candidate.exists():
        return candidate
    for ext in [".png", ".jpg", ".jpeg", ".pdf", ".eps", ".svg"]:
        c = root / (rel_path + ext)
        if c.exists():
            return c
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# Part 2: 图分类 & 知识包生成 (方案 C)
# ═══════════════════════════════════════════════════════════════════════════════

# ── 默认注册表 ────────────────────────────────────────────────────────────────

DEFAULT_REGISTRY = {
    "figure1_ict_stylized.png": {
        "paper_id": "Figure 1",
        "title": "The I-C-T Framework — Two Regimes of Attention Events",
        "type": "data",
        "generator": "generate_ict_figures.py::generate_figure1",
        "data_sources": [],
    },
    "figure2_covid_aal.png": {
        "paper_id": "Figure 2",
        "title": "COVID-19 AAL Four-Phase Case Study",
        "type": "data",
        "generator": "generate_ict_figures.py::generate_figure2",
        "data_sources": ["experiment/data/sp500-svi-covid.json", "data/processed/stock_returns_daily.parquet"],
    },
    "figure3_sector_diffusion.png": {
        "paper_id": "Figure 3",
        "title": "Sector Diffusion Ordering During COVID Phase 2",
        "type": "data",
        "generator": "generate_ict_figures.py::generate_figure3",
        "data_sources": ["experiment/data/sp500-svi-covid.json", "data/processed/stock_returns_daily.parquet"],
    },
    "causal-chain-3panel.png": {
        "paper_id": "Supplementary",
        "title": "Causal Chain — Three-Panel Diagram",
        "type": "conceptual",
        "generator": None,
        "data_sources": [],
    },
    "cognitive-diffusion-comparison.png": {
        "paper_id": "Supplementary",
        "title": "Cognitive Diffusion Comparison",
        "type": "conceptual",
        "generator": None,
        "data_sources": [],
    },
    "four-layer-diffusion.png": {
        "paper_id": "Supplementary",
        "title": "Four-Layer Diffusion Model",
        "type": "conceptual",
        "generator": None,
        "data_sources": [],
    },
}


def load_registry() -> dict:
    """加载注册表: 磁盘版 > 默认版"""
    if REGISTRY_PATH.exists():
        return json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))
    return dict(DEFAULT_REGISTRY)


def save_registry(registry: dict):
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    REGISTRY_PATH.write_text(json.dumps(registry, ensure_ascii=False, indent=2), encoding="utf-8")


def register_figure(filename: str, paper_id: str, fig_type: str,
                    generator: str = None, data_sources: list = None, title: str = ""):
    """注册一张新图到 registry"""
    registry = load_registry()
    registry[filename] = {
        "paper_id": paper_id,
        "title": title or filename,
        "type": fig_type,
        "generator": generator,
        "data_sources": data_sources or [],
    }
    save_registry(registry)
    print(f"✅ 已注册: {filename} → {paper_id} ({fig_type})")


# ── 代码提取 ──────────────────────────────────────────────────────────────────

def extract_generator_code(generator_ref: str, figures_dir: Path) -> str:
    if not generator_ref:
        return ""
    parts = generator_ref.split("::")
    if len(parts) != 2:
        return ""
    filename, funcname = parts
    filepath = figures_dir / filename
    if not filepath.exists():
        return f"(生成脚本不存在: {filepath})"

    code = filepath.read_text(encoding="utf-8")
    pattern = rf"(def {funcname}\(.*?\n(?:(?!^def ).*\n)*)"
    match = re.search(pattern, code, re.MULTILINE)
    if not match:
        return f"(函数 {funcname} 未找到)"

    func_code = match.group(1).strip()
    # 提取 docstring + 关键赋值行，不给全部绘图代码
    lines = func_code.split("\n")
    summary = []
    in_doc = False
    doc_done = False
    logic_count = 0

    for ln in lines:
        if '"""' in ln:
            if in_doc:
                in_doc = False
                doc_done = True
                summary.append(ln)
                continue
            else:
                in_doc = True
        if in_doc or not doc_done:
            summary.append(ln)
        elif doc_done:
            stripped = ln.strip()
            if stripped and not stripped.startswith("#"):
                # 保留关键逻辑行
                if any(kw in stripped for kw in [
                    "=", "return", "for ", "if ", "phases", "configs",
                    "color", "mid", "sigmoid", "ticker", "svi", "sector",
                    "pd.read", "json.load", "open(", "parquet",
                ]):
                    summary.append(ln)
                    logic_count += 1
                    if logic_count > 30:
                        summary.append("    # ... (truncated)")
                        break

    return "\n".join(summary)


# ── 论文上下文 ────────────────────────────────────────────────────────────────

def find_paper_context(paper_text: str, search_term: str) -> str:
    lines = paper_text.split("\n")
    contexts = []
    for i, line in enumerate(lines):
        if search_term.lower() in line.lower():
            start = max(0, i - 2)
            end = min(len(lines), i + 3)
            ctx = "\n".join(lines[start:end]).strip()
            if ctx and ctx not in contexts:
                contexts.append(ctx)
    return "\n\n---\n\n".join(contexts[:3]) if contexts else ""


# ── 主构建 ────────────────────────────────────────────────────────────────────

def build_pack(extracted: dict, figures_dir: Path) -> dict:
    """从解析结果 + 注册表构建知识包"""
    registry = load_registry()
    paper_text = extracted.get("text", "")

    # 建立已提取图 → 文件名的映射
    extracted_figs = {}
    for ef in extracted["figures"]:
        path = ef.get("extracted_image_path") or ef.get("source_path", "")
        if path:
            extracted_figs[Path(path).name] = ef

    figures = []
    needs_description = []

    # 遍历注册表中的图
    for filename, meta in registry.items():
        fig_path = figures_dir / filename
        ef = extracted_figs.pop(filename, None)

        entry = {
            "filename": filename,
            "paper_id": meta["paper_id"],
            "title": meta["title"],
            "type": meta["type"],
            "exists": fig_path.exists(),
            "file_size_kb": round(fig_path.stat().st_size / 1024, 1) if fig_path.exists() else 0,
            "caption": ef["caption"] if ef else "",
        }

        # 论文引用上下文
        if meta["paper_id"].startswith("Figure"):
            entry["paper_context"] = find_paper_context(paper_text, meta["paper_id"])
        else:
            entry["paper_context"] = ""

        if meta["type"] == "data":
            entry["generator_code"] = extract_generator_code(meta["generator"], figures_dir)
            entry["data_sources"] = [str(PROJECT_DIR / ds) for ds in meta["data_sources"]]
            entry["visual_description"] = ""
        else:
            entry["generator_code"] = ""
            entry["data_sources"] = []
            entry["visual_description"] = ""
            needs_description.append(filename)

        figures.append(entry)

    # 注册表外的图 (从论文/PDF 中提取到但未注册)
    for fname, ef in extracted_figs.items():
        figures.append({
            "filename": fname,
            "paper_id": "Unregistered",
            "title": ef.get("caption", fname),
            "type": "unknown",
            "exists": bool(ef.get("extracted_image_path")),
            "file_size_kb": 0,
            "caption": ef.get("caption", ""),
            "paper_context": "",
            "generator_code": "",
            "data_sources": [],
            "visual_description": "",
        })
        needs_description.append(fname)

    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "figures": figures,
        "tables": extracted["tables"],
        "needs_visual_description": needs_description,
        "stats": {
            "sections": len(extracted["sections"]),
            "tables": len(extracted["tables"]),
            "figures_total": len(figures),
            "figures_data": sum(1 for f in figures if f["type"] == "data"),
            "figures_conceptual": sum(1 for f in figures if f["type"] == "conceptual"),
            "figures_unregistered": sum(1 for f in figures if f["type"] == "unknown"),
        },
    }


# ── 描述填充 ──────────────────────────────────────────────────────────────────

def fill_descriptions(pack: dict, desc_path: Path) -> dict:
    if not desc_path.exists():
        print(f"⚠️ 描述文件不存在: {desc_path}")
        return pack
    descriptions = json.loads(desc_path.read_text(encoding="utf-8"))
    for fig in pack["figures"]:
        if fig["filename"] in descriptions:
            fig["visual_description"] = descriptions[fig["filename"]]
            if fig["filename"] in pack["needs_visual_description"]:
                pack["needs_visual_description"].remove(fig["filename"])
    return pack


# ── Markdown 输出 ─────────────────────────────────────────────────────────────

def format_md(pack: dict) -> str:
    lines = [
        "# Paper Assets — Figure & Table Descriptions",
        "",
        "> Auto-generated for Academic Review Board experts.",
        f"> Generated: {pack['generated_at']}",
        "",
        "---", "",
        "## Figures", "",
    ]

    for fig in pack["figures"]:
        type_label = {"data": "Data-driven", "conceptual": "Conceptual diagram"}.get(fig["type"], "Unregistered")
        lines.append(f"### {fig['paper_id']}: {fig['title']}")
        lines.append(f"- **File**: `{fig['filename']}` ({fig['file_size_kb']} KB)")
        lines.append(f"- **Type**: {type_label}")
        lines.append("")

        if fig["paper_context"]:
            lines.extend(["**Paper context:**", "```", fig["paper_context"], "```", ""])

        if fig["type"] == "data":
            if fig["data_sources"]:
                lines.append("**Data sources:**")
                for ds in fig["data_sources"]:
                    lines.append(f"- `{Path(ds).name}`")
                lines.append("")
            if fig["generator_code"]:
                lines.extend(["**Generating code:**", "```python", fig["generator_code"], "```", ""])

        if fig["visual_description"]:
            lines.extend(["**Visual description:**", fig["visual_description"], ""])
        elif fig["type"] in ("conceptual", "unknown"):
            lines.extend(["⚠️ *Visual description not yet generated.*", ""])

        lines.extend(["---", ""])

    if pack["tables"]:
        lines.extend(["## Tables", ""])
        for tab in pack["tables"]:
            lines.append(f"### {tab['id']}: {tab['title']}")
            md = tab.get("markdown", "")
            lines.append(md if md and md != "(inline in text)" else "*(inline in text)*")
            lines.extend(["", "---", ""])

    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="学术委员会审前一键资产预处理 (方案 C)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", nargs="?", help="论文文件 (.md / .pdf / .docx / .tex)")
    parser.add_argument("--figures-dir", default=None, help="图片目录 (默认: 项目 figures/)")
    parser.add_argument("--tex-root", default=None, help="LaTeX 项目根目录")
    parser.add_argument("--fill-descriptions", default=None, help="视觉描述 JSON 文件")
    parser.add_argument("--output-dir", default=str(CACHE_DIR), help="输出目录 (默认: cache/)")

    # 注册子命令
    parser.add_argument("--register", default=None, help="注册新图: 文件名")
    parser.add_argument("--paper-id", default=None, help="注册用: 论文中的 ID (如 'Figure 4')")
    parser.add_argument("--fig-type", choices=["data", "conceptual"], default=None, help="注册用: 图类型")
    parser.add_argument("--generator", default=None, help="注册用: 生成函数 (file.py::func)")
    parser.add_argument("--fig-title", default=None, help="注册用: 图标题")

    args = parser.parse_args()

    CACHE_DIR.mkdir(parents=True, exist_ok=True)

    # ── 注册模式 ──
    if args.register:
        if not args.paper_id or not args.fig_type:
            sys.exit("❌ 注册需要 --paper-id 和 --fig-type")
        register_figure(
            args.register, args.paper_id, args.fig_type,
            generator=args.generator, title=args.fig_title or "",
        )
        return

    # ── 正常模式 ──
    if not args.input:
        sys.exit("❌ 请指定论文文件路径 (.md / .pdf / .docx / .tex)")

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        sys.exit(f"❌ 文件不存在: {input_path}")

    figures_dir = Path(args.figures_dir) if args.figures_dir else DEFAULT_FIGURES_DIR
    output_dir = make_output_dir(args.output_dir)
    suffix = input_path.suffix.lower()

    print(f"📄 输入: {input_path.name} ({suffix})")
    print(f"📁 图片目录: {figures_dir}")
    print(f"📦 输出: {output_dir}")
    print()

    # Step 1: 解析论文
    if suffix == ".md":
        extracted = extract_markdown(str(input_path), figures_dir)
    elif suffix == ".pdf":
        extracted = extract_pdf(str(input_path), output_dir)
    elif suffix == ".docx":
        extracted = extract_docx(str(input_path), output_dir)
    elif suffix in (".tex", ".latex"):
        extracted = extract_tex(str(input_path), output_dir, tex_root=args.tex_root)
    else:
        sys.exit(f"❌ 不支持的格式: {suffix}")

    print(f"  解析完成: {len(extracted['sections'])} 章节, "
          f"{len(extracted['tables'])} 表格, {len(extracted['figures'])} 图片")

    # Step 2: 构建知识包
    pack = build_pack(extracted, figures_dir)

    # Step 3: 填充描述
    if args.fill_descriptions:
        desc_path = Path(args.fill_descriptions).expanduser().resolve()
        pack = fill_descriptions(pack, desc_path)
        filled = sum(1 for f in pack["figures"] if f["visual_description"])
        print(f"  填充了 {filled} 个视觉描述")

    # 写 JSON
    json_path = output_dir / "paper-figure-pack.json"
    json_path.write_text(json.dumps(pack, ensure_ascii=False, indent=2), encoding="utf-8")

    # 写 Markdown
    md_path = output_dir / "paper-figure-pack.md"
    md_path.write_text(format_md(pack), encoding="utf-8")

    # 统计
    s = pack["stats"]
    print(f"\n✅ 完成!")
    print(f"   JSON → {json_path}")
    print(f"   Markdown → {md_path}")
    print(f"   📊 {s['figures_data']} 数据图, {s['figures_conceptual']} 概念图, "
          f"{s['figures_unregistered']} 未注册, {s['tables']} 表格")

    needs = pack["needs_visual_description"]
    if needs:
        print(f"\n⚠️  {len(needs)} 张图需要视觉描述:")
        for fn in needs:
            print(f"   - {fn}")
        print(f"\n   用 image 工具生成描述后:")
        print(f"   python3 {Path(__file__).name} {input_path.name} "
              f"--fill-descriptions figure-descriptions.json")


if __name__ == "__main__":
    main()
